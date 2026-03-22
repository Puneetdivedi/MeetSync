import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF

def ensure_outputs_dir():
    if not os.path.exists("outputs"):
        os.makedirs("outputs")

def add_docx_table(doc, items, headers, keys):
    if not items:
        doc.add_paragraph("None identified.")
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        
    # Rows
    for item in items:
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            row_cells[i].text = str(item.get(key, ""))

def generate_docx_content(doc, title, results):
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = 1 # Center
    
    # Meta data
    meta = doc.add_paragraph()
    meta.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta.alignment = 1
    
    doc.add_page_break()
    
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(results.get("summary", ""))
    
    doc.add_heading("Action Items", level=1)
    if "action_items_raw" in results and "items" in results["action_items_raw"]:
        add_docx_table(doc, results["action_items_raw"]["items"], ["Assignee", "Task", "Deadline"], ["assignee", "task", "deadline"])
    else:
        doc.add_paragraph(results.get("action_items", ""))
    
    doc.add_heading("Key Decisions", level=1)
    if "decisions_raw" in results and "decisions" in results["decisions_raw"]:
        add_docx_table(doc, results["decisions_raw"]["decisions"], ["Decision", "Context"], ["decision", "context"])
    else:
        doc.add_paragraph(results.get("decisions", ""))
    
    doc.add_heading("Follow-up Email Draft", level=1)
    doc.add_paragraph(results.get("email_draft", ""))
    
    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph(results.get("transcript", ""))

def export_docx(results, filename="MeetSync_Enterprise_Report.docx"):
    ensure_outputs_dir()
    filepath = os.path.join("outputs", filename)
    doc = Document()
    generate_docx_content(doc, "MeetSync Enterprise Intelligence Report", results)
    doc.save(filepath)
    return filepath

class EnterprisePDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.set_text_color(15, 23, 42) # Slate-900
        self.cell(0, 10, 'MeetSync Enterprise Intelligence Report', 0, 1, 'C')
        self.set_font('Arial', 'I', 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', 0, 1, 'C')
        self.ln(10)
        
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def add_pdf_table(pdf, items, headers, keys, col_widths):
    if not items:
        pdf.set_font("Arial", 'I', 11)
        pdf.cell(0, 8, "None identified.", 0, 1)
        return
        
    # Header
    pdf.set_font("Arial", 'B', 10)
    pdf.set_fill_color(241, 245, 249) # Slate-100
    for i, header in enumerate(headers):
        pdf.cell(col_widths[i], 8, header, 1, 0, 'L', True)
    pdf.ln()
    
    # Rows
    pdf.set_font("Arial", '', 10)
    for item in items:
        row_texts = [str(item.get(k, "")) for k in keys]
        for i, txt in enumerate(row_texts):
            safe_txt = txt.encode('latin-1', 'replace').decode('latin-1')
            # Truncate string to roughly fit the FPDF cell width without overflowing awkwardly
            char_limit = int(col_widths[i] / 2.5) 
            if len(safe_txt) > char_limit:
                safe_txt = safe_txt[:char_limit-3] + "..."
            pdf.cell(col_widths[i], 8, safe_txt, 1, 0, 'L')
        pdf.ln()

def export_pdf(results, filename="MeetSync_Enterprise_Report.pdf"):
    ensure_outputs_dir()
    filepath = os.path.join("outputs", filename)
    pdf = EnterprisePDF()
    pdf.add_page()
    
    def safe_print_section(title, content):
        pdf.set_font("Arial", 'B', 14)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(0, 10, title, 0, 1)
        pdf.set_font("Arial", size=11)
        pdf.set_text_color(50, 50, 50)
        safe_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, safe_content)
        pdf.ln(5)

    # Summary
    safe_print_section("Executive Summary", results.get("summary", ""))

    # Action Items
    pdf.set_font("Arial", 'B', 14)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 10, "Action Items", 0, 1)
    if "action_items_raw" in results and "items" in results["action_items_raw"]:
        add_pdf_table(pdf, results["action_items_raw"]["items"], 
                      ["Assignee", "Task", "Deadline"], 
                      ["assignee", "task", "deadline"], 
                      [40, 120, 30])
    else:
        safe_print_section("", results.get("action_items", ""))
    pdf.ln(5)

    # Decisions
    pdf.set_font("Arial", 'B', 14)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 10, "Key Decisions", 0, 1)
    if "decisions_raw" in results and "decisions" in results["decisions_raw"]:
        add_pdf_table(pdf, results["decisions_raw"]["decisions"], 
                      ["Decision", "Context"], 
                      ["decision", "context"], 
                      [95, 95])
    else:
        safe_print_section("", results.get("decisions", ""))
    pdf.ln(5)

    # Email
    safe_print_section("Follow-up Email Draft", results.get("email_draft", ""))
    
    # Transcript
    pdf.add_page()
    safe_print_section("Full Transcript", results.get("transcript", ""))
        
    pdf.output(filepath)
    return filepath
