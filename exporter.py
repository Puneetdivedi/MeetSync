import os
from docx import Document
from fpdf import FPDF

def ensure_outputs_dir():
    if not os.path.exists("outputs"):
        os.makedirs("outputs")

def generate_docx_content(doc, title, results):
    doc.add_heading(title, 0)
    
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(results.get("summary", ""))
    
    doc.add_heading("Action Items", level=1)
    doc.add_paragraph(results.get("action_items", ""))
    
    doc.add_heading("Key Decisions", level=1)
    doc.add_paragraph(results.get("decisions", ""))
    
    doc.add_heading("Follow-up Email Draft", level=1)
    doc.add_paragraph(results.get("email_draft", ""))
    
    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph(results.get("transcript", ""))

def export_docx(results, filename="MeetSync_Report.docx"):
    ensure_outputs_dir()
    filepath = os.path.join("outputs", filename)
    doc = Document()
    generate_docx_content(doc, "Meeting Intelligence Report", results)
    doc.save(filepath)
    return filepath

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'MeetSync Meeting Report', 0, 1, 'C')
        self.ln(10)

def export_pdf(results, filename="MeetSync_Report.pdf"):
    ensure_outputs_dir()
    filepath = os.path.join("outputs", filename)
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    
    sections = [
        ("Executive Summary", results.get("summary", "")),
        ("Action Items", results.get("action_items", "")),
        ("Key Decisions", results.get("decisions", "")),
        ("Follow-up Email Draft", results.get("email_draft", "")),
        ("Full Transcript", results.get("transcript", ""))
    ]
    
    for title, content in sections:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, title, 0, 1)
        pdf.set_font("Arial", size=11)
        # Handle utf-8 to latin-1 encoding issues natively in FPDF by using multi_cell
        safe_content = content.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 6, safe_content)
        pdf.ln(5)
        
    pdf.output(filepath)
    return filepath
